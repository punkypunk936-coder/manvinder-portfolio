from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "assets" / "manvinder-arora-360-resume.docx"

INK = RGBColor(20, 25, 29)
MUTED = RGBColor(82, 92, 100)
ACCENT = RGBColor(8, 123, 104)
FONT = "Arial"


def set_font(run, size=8.8, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_hyperlink(paragraph, text, url, color=ACCENT, bold=False, size=8.6):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(fonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), str(color))
    r_pr.append(color_el)
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(size_el)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bottom_border(paragraph, color="087B68", size="10", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_cell_shading(cell, fill="F2F5F4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def mark_row_as_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(8.8)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.04

    heading = doc.styles["Heading 1"]
    heading.font.name = FONT
    heading._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    heading.font.size = Pt(9.4)
    heading.font.bold = True
    heading.font.color.rgb = ACCENT
    heading.paragraph_format.space_before = Pt(5.5)
    heading.paragraph_format.space_after = Pt(2.5)

    bullets = doc.styles["List Bullet"]
    bullets.font.name = FONT
    bullets._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    bullets._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    bullets.font.size = Pt(8.55)


def section_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    set_font(run, size=9.4, bold=True, color=ACCENT)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.11)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.6)
    p.paragraph_format.line_spacing = 1.03
    p.paragraph_format.keep_together = True
    set_font(p.add_run(text), size=8.55)
    return p


def role_header(doc, company, company_url, title, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2.5)
    p.paragraph_format.space_after = Pt(1.6)
    p.paragraph_format.keep_with_next = True
    add_hyperlink(p, company, company_url, color=INK, bold=True, size=9.35)
    set_font(p.add_run(f" | {title}"), size=9.35, bold=True)
    set_font(p.add_run(f"    {dates}"), size=8.1, bold=True, color=MUTED)


def linked_list_line(doc, label, links):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.03
    set_font(p.add_run(f"{label}: "), size=8.55, bold=True)
    for index, (name, url) in enumerate(links):
        if index:
            set_font(p.add_run("  |  "), size=8.2, color=MUTED)
        add_hyperlink(p, name, url, bold=True, size=8.45)


def build_resume():
    doc = Document()
    doc.core_properties.title = "Manvinder Arora - Crypto-Native Operator and AI Builder"
    doc.core_properties.subject = "One-page professional resume"
    doc.core_properties.author = "Manvinder Arora"
    doc.core_properties.keywords = "crypto, DeFi, growth, marketing, community, business development, trading, AI"

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.38)
    section.bottom_margin = Inches(0.34)
    section.left_margin = Inches(0.52)
    section.right_margin = Inches(0.52)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.18)
    configure_styles(doc)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    add_hyperlink(
        footer,
        "Portfolio: punkypunk936-coder.github.io/manvinder-portfolio",
        "https://punkypunk936-coder.github.io/manvinder-portfolio/",
        color=MUTED,
        size=7.4,
    )

    name = doc.add_paragraph()
    name.paragraph_format.space_before = Pt(0)
    name.paragraph_format.space_after = Pt(0)
    set_font(name.add_run("MANVINDER ARORA"), size=20.5, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(2.5)
    set_font(title.add_run("CRYPTO-NATIVE OPERATOR & AI BUILDER"), size=8.9, bold=True, color=ACCENT)

    contact = doc.add_paragraph()
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(4)
    set_font(contact.add_run("India (UTC+5:30)  |  "), size=8, color=MUTED)
    add_hyperlink(contact, "Email", "mailto:mani.arora03@gmail.com", color=MUTED, size=8)
    set_font(contact.add_run("  |  "), size=8, color=MUTED)
    add_hyperlink(contact, "X", "https://x.com/0xgoodie", color=MUTED, size=8)
    set_font(contact.add_run("  |  "), size=8, color=MUTED)
    add_hyperlink(contact, "LinkedIn", "https://www.linkedin.com/in/manvinderarora/", color=MUTED, size=8)
    set_font(contact.add_run("  |  "), size=8, color=MUTED)
    add_hyperlink(contact, "Substack", "https://manvinder.substack.com", color=MUTED, size=8)
    set_font(contact.add_run("  |  "), size=8, color=MUTED)
    add_hyperlink(contact, "Portfolio", "https://punkypunk936-coder.github.io/manvinder-portfolio/", color=MUTED, size=8)
    add_bottom_border(contact)

    summary = doc.add_paragraph()
    summary.paragraph_format.space_before = Pt(3)
    summary.paragraph_format.space_after = Pt(3.5)
    summary.paragraph_format.line_spacing = 1.06
    set_font(
        summary.add_run(
            "Crypto-native operator with five years across community, content, growth, marketing and business development. I helped Timeswap grow TVL by about 50%, brought in repeat liquidity providers with roughly $290K in starting capital, helped scale large communities, and later supported FusionX launches tied to $100M+ in 30-day volume. Outside work, I trade onchain, write about markets and build small AI tools that make my own workflows easier."
        ),
        size=8.75,
    )

    metrics = [
        ("~50%", "TVL growth"),
        ("~$290K", "LP capital"),
        ("$100M+", "30-day volume"),
        ("150K+", "brand followers"),
        ("50K + 50K", "Discord + X"),
        ("2,000+", "campaign users"),
    ]
    table = doc.add_table(rows=1, cols=6)
    table.autofit = False
    table.allow_autofit = False
    mark_row_as_header(table.rows[0])
    for cell, (value, label) in zip(table.rows[0].cells, metrics):
        cell.width = Inches(1.24)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2.5)
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing = 1.0
        set_font(p.add_run(f"{value}\n"), size=9.6, bold=True, color=ACCENT)
        set_font(p.add_run(label), size=6.9, color=MUTED)

    section_heading(doc, "Experience")
    role_header(doc, "FusionX Finance", "https://docs.fusionx.finance/", "Marketing & BD Head", "Jul 2024 - Mar 2026")
    bullet(doc, "Ran marketing and BD across EVM and SVM products, working with founders, chain teams, creators and communities on launches and partner campaigns.")
    bullet(doc, "Helped grow SummitX to 150K+ followers and supported CobaltX campaigns tied to more than $100M in trading volume over 30 days.")
    bullet(doc, "Built direct feedback loops with traders, whales and LPs; turned user friction into clearer content, onboarding and product feedback.")

    role_header(doc, "Timeswap", "https://timeswap.io/", "Marketing & Community Head", "Jul 2021 - Jun 2024")
    bullet(doc, "Joined through community and product education, then took on content, growth, campaigns and LP relationships as the protocol grew.")
    bullet(doc, "Helped support roughly 50% TVL growth and personally onboarded repeat LPs with about $290K in initial capital ($200K, $60K and $30K).")
    bullet(doc, "Grew and moderated a 50K+ Discord, helped scale X to roughly 50K followers, and ran campaigns that added 2,000+ active users (about 15%) in one month.")
    bullet(doc, "Worked directly with creators and community leaders; wrote threads, FAQs and launch messages that made oracle-less lending easier to understand.")

    role_header(doc, "Independent", "https://x.com/0xgoodie", "Onchain Trader, Writer & AI Builder", "2021 - Present")
    bullet(doc, "Use Hyperliquid, Lighter, Variational, Nado and Aster as a real trader; received a top-100 Lighter airdrop allocation through sustained product usage.")
    bullet(doc, "Trade crypto, onchain equities, indices and commodities; compare venues through execution, liquidity, fees, funding, incentives and the full user journey.")
    bullet(doc, "Write about perp DEXs, stablecoins, tokenized markets and community work, and build personal tools when an existing workflow feels clumsy.")

    section_heading(doc, "Linked Work")
    linked_list_line(
        doc,
        "Built for my own use",
        [
            ("Meme Scout Agent", "https://github.com/punkypunk936-coder/meme-scout-agent"),
            ("Crypto Trading Agent", "https://github.com/punkypunk936-coder/crypto-trading-agent"),
            ("X-to-Substack", "https://github.com/punkypunk936-coder/x-to-substack-pipeline"),
        ],
    )
    linked_list_line(
        doc,
        "Selected writing",
        [
            ("Perp DEXs: Crypto's clearest PMF", "https://manvinder.substack.com/p/perp-dexs-cryptos-clearest-pmf"),
            ("TradFi became crypto", "https://manvinder.substack.com/p/crypto-didnt-become-tradfi-tradfi"),
            ("Community manager", "https://x.com/0xgoodie/status/2075661908217364766"),
        ],
    )

    section_heading(doc, "What I Bring")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.04
    set_font(p.add_run("Growth, community, content, GTM and partnerships; liquidity-provider, whale and creator relationships; DeFi and perp DEX fluency; product feedback; clear writing; AI-assisted workflow design."), size=8.5)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_resume()
